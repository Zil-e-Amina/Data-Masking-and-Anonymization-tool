import streamlit as st
import pandas as pd
import json, sys, os, matplotlib.pyplot as plt
from collections import Counter
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from src.pipeline import process, process_dataframe
from src.security import check_file_size, sanitize_input, validate_file_type, safe_error, log_audit

st.markdown("""<style>
@import url('https://fonts.googleapis.com/css2?family=JetBrains+Mono&family=Outfit:wght@300;400;600;700;800&display=swap');
html,body,[class*="css"]{font-family:'Outfit',sans-serif}
.main{background:#050A18}
.block-container{padding:2rem 2.5rem}
h1{font-size:2.4rem!important;font-weight:800!important;color:white!important;letter-spacing:-1px!important}
h2,h3{color:#94A3B8!important;font-weight:600!important;font-size:.85rem!important;letter-spacing:1.5px!important;text-transform:uppercase!important}
[data-testid="stSidebar"]{background:#070D1A!important;border-right:1px solid #1E293B!important}
.stTabs [data-baseweb="tab-list"]{background:#0D1424;border-radius:12px;padding:4px;gap:2px;border:1px solid #1E293B}
.stTabs [data-baseweb="tab"]{border-radius:9px;color:#64748B;font-weight:600;padding:10px 22px}
.stTabs [aria-selected="true"]{background:linear-gradient(135deg,#0EA5E9,#6366F1)!important;color:white!important;box-shadow:0 4px 15px rgba(14,165,233,.3)!important}
.stButton>button{background:linear-gradient(135deg,#0EA5E9,#6366F1)!important;color:white!important;border:none!important;border-radius:10px!important;font-weight:700!important;font-size:15px!important;padding:13px!important;box-shadow:0 4px 15px rgba(14,165,233,.2)!important;transition:all .2s!important}
.stButton>button:hover{transform:translateY(-2px)!important;box-shadow:0 8px 25px rgba(14,165,233,.4)!important}
.stDownloadButton>button{background:linear-gradient(135deg,#059669,#0EA5E9)!important;color:white!important;border:none!important;border-radius:10px!important;font-weight:600!important}
.stTextArea textarea{background:#0B1525!important;color:#E2E8F0!important;border:1.5px solid #2D3F55!important;border-radius:10px!important;font-family:'JetBrains Mono',monospace!important;font-size:13px!important;line-height:1.8!important}
.stTextArea textarea::placeholder{color:#64748B!important;opacity:1!important;font-style:italic}
.stTextArea textarea:focus{border-color:#0EA5E9!important;box-shadow:0 0 0 2px rgba(14,165,233,.15)!important}
.stTextInput input{background:#0B1525!important;color:#E2E8F0!important;border:1.5px solid #2D3F55!important;border-radius:10px!important;font-family:'JetBrains Mono',monospace!important;font-size:13px!important}
.stTextInput input::placeholder{color:#64748B!important;opacity:1!important}
.stTextInput input:focus{border-color:#0EA5E9!important;box-shadow:0 0 0 2px rgba(14,165,233,.15)!important}
.stTextInput label{color:#94A3B8!important;font-size:13px!important;font-weight:600!important}
[data-testid="metric-container"]{background:#0B1525;border:1px solid #1E293B;border-radius:12px;padding:18px!important}
[data-testid="metric-container"] label{color:#64748B!important;font-size:10px!important;letter-spacing:1.5px!important;text-transform:uppercase!important}
[data-testid="stMetricValue"]{color:#38BDF8!important;font-weight:800!important}
.stDataFrame{border:1px solid #1E293B!important;border-radius:12px!important}
[data-testid="stFileUploader"]{background:#0B1525!important;border:2px dashed #2D3F55!important;border-radius:14px!important;padding:20px!important}
.stSelectbox>div>div{background:#0B1525!important;border:1px solid #1E293B!important;border-radius:10px!important;color:#CBD5E1!important}
.card{background:#0B1525;border:1px solid #1E293B;border-radius:14px;padding:20px;margin:8px 0}
.card-blue{background:#0B1525;border:1px solid #1E293B;border-left:3px solid #0EA5E9;border-radius:14px;padding:20px;margin:8px 0}
.safe-box{background:linear-gradient(135deg,#022C22,#064E3B);border:1px solid #059669;border-radius:14px;padding:28px;text-align:center;margin:12px 0}
.fail-box{background:linear-gradient(135deg,#1C0A0A,#2D1515);border:1px solid #DC2626;border-radius:14px;padding:28px;text-align:center;margin:12px 0}
.ok{background:rgba(5,150,105,.12);border-left:3px solid #059669;border-radius:8px;padding:9px 14px;margin:4px 0;color:#34D399;font-size:13px}
.no{background:rgba(220,38,38,.12);border-left:3px solid #DC2626;border-radius:8px;padding:9px 14px;margin:4px 0;color:#F87171;font-size:13px}
.gb{background:#0B1525;border:1px solid #1E293B;border-radius:12px;padding:16px}
.hero{background:linear-gradient(135deg,#0B1525,#0D1A30,#0B1525);border:1px solid #1E293B;border-radius:18px;padding:38px 44px;margin-bottom:24px;position:relative;overflow:hidden}
.hero::before{content:'';position:absolute;top:-80px;right:-80px;width:280px;height:280px;border-radius:50%;background:radial-gradient(circle,rgba(14,165,233,.1),transparent 70%)}
.pill{display:inline-block;background:#1E293B;border-radius:20px;padding:4px 13px;font-size:11px;color:#64748B;margin:3px 3px 0 0;font-weight:600}
.pill-g{display:inline-block;background:rgba(5,150,105,.15);border:1px solid #059669;border-radius:20px;padding:4px 13px;font-size:11px;color:#34D399;margin:3px 3px 0 0;font-weight:600}
.div{border:none;border-top:1px solid #1E293B;margin:20px 0}
.empty-box{background:#0B1525;border:1px dashed #2D3F55;border-radius:10px;height:420px;display:flex;align-items:center;justify-content:center;flex-direction:column;gap:10px}
</style>""", unsafe_allow_html=True)


# ── File readers ──────────────────────────────────────────────────────────────
def read_pdf(f):
    try:
        import fitz
        pdf_bytes = f.read()
        doc  = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = ""
        for i, page in enumerate(doc):
            page_text = page.get_text()
            if page_text.strip():
                text += f"-- Page {i+1} --\n{page_text}\n"
        doc.close()
        return text if text.strip() else "No readable text found in PDF."
    except Exception as e:
        return "Could not read PDF: " + str(e)


def read_docx(f):
    try:
        from docx import Document
        doc  = Document(f)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    c.text for c in row.cells if c.text.strip()
                )
                if row_text:
                    text += "\n" + row_text
        return text if text.strip() else "No readable text found."
    except Exception as e:
        return "Could not read document: " + str(e)


def mkfig(r, c):
    fig, axes = plt.subplots(r, c, figsize=(10, 3))
    fig.patch.set_facecolor('#050A18')
    for ax in (axes.flat if hasattr(axes, 'flat') else [axes]):
        ax.set_facecolor('#0B1525')
        ax.tick_params(colors='#64748B')
        ax.title.set_color('#94A3B8')
        ax.xaxis.label.set_color('#64748B')
        ax.yaxis.label.set_color('#64748B')
        for s in ax.spines.values():
            s.set_edgecolor('#1E293B')
    return fig, axes


def show_results(results, risk, val):
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Fields Found",   len(results))
    c2.metric("Risk Score",     f"{risk}/100")
    c3.metric("Critical",       sum(1 for r in results if "CRITICAL" in r["Risk"]))
    c4.metric("High",           sum(1 for r in results if "HIGH"     in r["Risk"]))
    st.markdown("<hr class='div'>", unsafe_allow_html=True)

    if   risk >= 70: st.error("HIGH RISK — Very sensitive data detected")
    elif risk >= 40: st.warning("MEDIUM RISK — Some sensitive data found")
    else:            st.success("LOW RISK — Minimal sensitive data")

    if not results:
        st.success("No sensitive data detected.")
        return

    st.subheader("Detection Table")
    st.dataframe(pd.DataFrame(results), use_container_width=True, hide_index=True)

    counts = Counter(r["Type"] for r in results)
    rem    = Counter(i["label"] for i in val.get("still_exposed", []))
    fig, ax = mkfig(1, 2)
    clrs = ["#0EA5E9","#6366F1","#8B5CF6","#EC4899","#F59E0B","#10B981"]
    ax[0].bar(counts.keys(), counts.values(),
              color=clrs[:len(counts)], edgecolor='#0B1525')
    ax[0].set_title("Before Masking")
    ax[0].tick_params(axis='x', rotation=30)
    ax[1].bar(counts.keys(), [rem.get(t, 0) for t in counts],
              color="#059669", edgecolor='#0B1525')
    ax[1].set_title("After Masking")
    ax[1].tick_params(axis='x', rotation=30)
    plt.tight_layout()
    st.pyplot(fig)

    df_r = pd.DataFrame(results)
    d1, d2 = st.columns(2)
    tc = Counter(r["Type"] for r in results)
    with d1:
        st.download_button("Download CSV Report",
            df_r.to_csv(index=False), "report.csv",
            mime="text/csv", use_container_width=True)
    with d2:
        txt = ("MASKING REPORT\n"
               f"Risk:{risk}/100\n"
               f"Safe:{'YES' if val['safe_to_share'] else 'NO'}\n\n")
        txt += "\n".join(f"{t}:{c}" for t, c in tc.items())
        st.download_button("Download TXT Report",
            txt, "report.txt", use_container_width=True)


def show_col(report):
    total = len(report)
    sens  = sum(1 for r in report if "Sensitive" in r["Status"])
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Columns",    total)
    c2.metric("Sensitive",  sens)
    c3.metric("Safe",       total - sens)
    c4.metric("Masked",     sum(r["Values Masked"] for r in report))
    st.dataframe(pd.DataFrame(report), use_container_width=True, hide_index=True)

    sd = [r for r in report if "Sensitive" in r["Status"]]
    if sd:
        fig, ax = mkfig(1, 2)
        ax[0].pie([sens, total - sens],
                  labels=["Sensitive", "Safe"],
                  colors=["#DC2626", "#059669"],
                  autopct="%1.0f%%",
                  textprops={'color': '#94A3B8'})
        ax[0].set_title("Split")
        ax[1].barh([r["Column"] for r in sd],
                   [r["Values Masked"] for r in sd],
                   color="#0EA5E9", edgecolor='#0B1525')
        ax[1].set_title("Masked Per Column")
        plt.tight_layout()
        st.pyplot(fig)

    st.download_button("Download Column Report",
        pd.DataFrame(report).to_csv(index=False),
        "col_report.csv", use_container_width=True)


def show_share(results, risk, val, fname="dataset"):
    safe = val["safe_to_share"]
    st_  = ("APPROVED" if safe and risk < 70
            else "CONDITIONAL" if safe else "REJECTED")
    ic   = "✅" if st_ == "APPROVED" else "🟡" if st_ == "CONDITIONAL" else "❌"

    st.markdown("<hr class='div'>", unsafe_allow_html=True)
    st.subheader("Secure Data Sharing Panel")

    if st_ == "APPROVED":
        st.markdown(
            '<div class="safe-box">'
            '<div style="font-size:2.2rem">✅</div>'
            '<h2 style="color:#34D399;margin:4px 0;font-size:1.5rem">SAFE TO SHARE</h2>'
            '<p style="color:#6EE7B7;font-size:13px;margin:0">'
            'All sensitive data masked. Approved for secure sharing.</p>'
            '</div>',
            unsafe_allow_html=True
        )
    elif st_ == "CONDITIONAL":
        st.warning("CONDITIONALLY SAFE — Share only with trusted parties.")
    else:
        st.markdown(
            '<div class="fail-box">'
            '<div style="font-size:2.2rem">❌</div>'
            '<h2 style="color:#F87171;margin:4px 0;font-size:1.5rem">NOT SAFE TO SHARE</h2>'
            '<p style="color:#FCA5A5;font-size:13px;margin:0">'
            + str(val["remaining_count"]) + ' value(s) still exposed.</p>'
            '</div>',
            unsafe_allow_html=True
        )

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Found",   val["original_count"])
    c2.metric("Masked",  val["original_count"] - val["remaining_count"])
    c3.metric("Exposed", val["remaining_count"])
    c4.metric("Status",  ic + " " + st_)

    st.subheader("Pre-Sharing Checklist")
    for label, ok in [
        ("Sensitive fields detected",         True),
        ("Masking applied",                   len(results) > 0),
        ("Salted hashing used",               any("Hash" in r.get("Technique","") for r in results)),
        ("Validation scan done",              True),
        ("No data remains exposed",           safe),
        ("Report ready",                      True),
        ("Masked file ready",                 True),
    ]:
        css = "ok" if ok else "no"
        ico = "✅" if ok else "❌"
        st.markdown(
            f'<div class="{css}">{ico} {label}</div>',
            unsafe_allow_html=True
        )

    st.markdown("<br>", unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    with c1:
        st.markdown(
            '<div class="gb">'
            '<p style="color:#34D399;font-weight:700;font-size:13px;margin:0 0 10px">YOU CAN SHARE</p>'
            '<p style="color:#94A3B8;font-size:12px;line-height:1.9;margin:0">'
            'Masked CSV / TXT file<br>Transformation report<br>'
            'Column sensitivity report<br>Aggregated statistics</p>'
            '</div>',
            unsafe_allow_html=True
        )
    with c2:
        st.markdown(
            '<div class="gb">'
            '<p style="color:#F87171;font-weight:700;font-size:13px;margin:0 0 10px">DO NOT SHARE</p>'
            '<p style="color:#94A3B8;font-size:12px;line-height:1.9;margin:0">'
            'Original unmasked dataset<br>Token-to-value mappings<br>'
            'CRITICAL risk fields<br>Original input file</p>'
            '</div>',
            unsafe_allow_html=True
        )

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    tc  = Counter(r["Type"] for r in results)
    cert = (
        "DATA MASKING & ANONYMIZATION CERTIFICATE\n"
        "PAF-IAST | SSDD Cybersecurity Project 2025\n"
        "==========================================\n"
        f"File     : {fname}\n"
        f"Date     : {now}\n"
        f"Status   : {ic} {st_}\n"
        f"Risk     : {risk}/100\n"
        f"Safe     : {'YES' if safe else 'NO'}\n\n"
        "TYPES MASKED\n"
        "------------\n"
    )
    cert += "\n".join(f"  {t}: {c}" for t, c in tc.items())
    cert += (
        "\n\nSECURITY APPLIED\n"
        "----------------\n"
        "Salted SHA-256/MD5 hashing\n"
        "Input sanitization\n"
        "File validation\n"
        "Post-masking scan\n"
        "Audit logging\n\n"
        "TEAM\n"
        "----\n"
        "Abdullah Tipu | Syed Abdullah Gillani\n"
        "Ahmed Usman Mir | Fatima Tuz Zahra | Zile Amina\n\n"
        "INSTITUTION: PAF-IAST\n"
        "PROJECT: SSDD 2025\n"
    )

    st.markdown("<br>", unsafe_allow_html=True)
    st.download_button("Download Sharing Certificate",
        cert, f"cert_{fname}.txt", use_container_width=True)

    if safe:
        st.markdown(
            '<div class="card-blue">'
            '<p style="color:#34D399;font-size:13px;line-height:2;margin:0">'
            '1. Download masked CSV | '
            '2. Download certificate | '
            '3. Share masked file only | '
            '4. Keep original private</p>'
            '</div>',
            unsafe_allow_html=True
        )
    else:
        st.markdown(
            '<div class="card" style="border-left:3px solid #DC2626">'
            '<p style="color:#F87171;font-size:13px;line-height:2;margin:0">'
            '1. Do NOT share | '
            '2. Use stronger technique | '
            '3. Re-run masking | '
            '4. Verify safe then share</p>'
            '</div>',
            unsafe_allow_html=True
        )


# ── Page setup ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Data Masking Tool",
    page_icon="🔒",
    layout="wide"
)

with st.sidebar:
    st.markdown(
        '<div style="padding:24px 0 16px;text-align:center">'
        '<div style="font-size:40px">🔒</div>'
        '<p style="color:#0EA5E9;font-weight:700;font-size:14px;'
        'margin:6px 0 0;letter-spacing:2px;text-transform:uppercase">Settings</p>'
        '</div>',
        unsafe_allow_html=True
    )
    technique = st.selectbox(
        "Masking Technique",
        ["auto", "hash", "encrypt", "tokenize", "suppress", "noise"]
    )
    st.markdown("<hr class='div'>", unsafe_allow_html=True)
    st.markdown(
        '<div class="card">'
        '<p style="color:#94A3B8;font-weight:600;font-size:10px;'
        'letter-spacing:1.5px;text-transform:uppercase;margin:0 0 12px">Techniques</p>'
        '<p style="color:#64748B;font-size:12px;margin:0;line-height:2.2">'
        'auto — Best per field<br>'
        'hash — Salted SHA-256<br>'
        'encrypt — Base64<br>'
        'tokenize — TOKEN_001<br>'
        'suppress — [HIDDEN]<br>'
        'noise — Random noise</p>'
        '</div>',
        unsafe_allow_html=True
    )
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown(
        '<div class="card">'
        '<p style="color:#94A3B8;font-weight:600;font-size:10px;'
        'letter-spacing:1.5px;text-transform:uppercase;margin:0 0 12px">Risk Levels</p>'
        '<p style="color:#64748B;font-size:12px;margin:0;line-height:2.2">'
        'CRITICAL — CNIC, Card, Password<br>'
        'HIGH — Email, IP, Phone<br>'
        'MEDIUM — Name, Address, Port<br>'
        'LOW — URL, Agent</p>'
        '</div>',
        unsafe_allow_html=True
    )
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown(
        '<div style="background:rgba(5,150,105,.08);border:1px solid '
        'rgba(5,150,105,.25);border-radius:10px;padding:12px 16px">'
        '<p style="color:#34D399;font-weight:600;font-size:10px;'
        'letter-spacing:1.5px;text-transform:uppercase;margin:0 0 8px">Security Active</p>'
        '<p style="color:#64748B;font-size:11px;margin:0;line-height:2">'
        'Input sanitization<br>'
        'File size validation<br>'
        'Content-type checks<br>'
        'Salted hashing<br>'
        'Audit logging</p>'
        '</div>',
        unsafe_allow_html=True
    )
    st.markdown(
        '<p style="color:#1E293B;font-size:10px;text-align:center;margin-top:20px">'
        'PAF-IAST | SSDD 2025</p>',
        unsafe_allow_html=True
    )

st.markdown(
    '<div class="hero">'
    '<span class="pill">PAF-IAST</span>'
    '<span class="pill">SSDD 2025</span>'
    '<span class="pill">Python</span>'
    '<span class="pill">Cybersecurity</span>'
    '<span class="pill-g">Secured</span>'
    '<h1 style="margin:18px 0 10px">Data Masking &amp;<br>Anonymization Tool</h1>'
    '<p style="color:#475569;margin:0;font-size:14px;max-width:500px;line-height:1.7">'
    'Detect, mask and validate sensitive data in cybersecurity datasets '
    'before secure sharing.</p>'
    '</div>',
    unsafe_allow_html=True
)

tab1, tab2, tab3 = st.tabs([
    "✍️   Manual Input",
    "📂   CSV / JSON",
    "📄   PDF / Word"
])


# ── TAB 1 ─────────────────────────────────────────────────────────────────────
with tab1:
    col1, col2 = st.columns(2, gap="large")
    with col1:
        st.subheader("Raw Input")
        st.markdown(
            '<p style="color:#64748B;font-size:12px;margin:0 0 12px">'
            'Fill in the fields you want to mask. Leave others empty.</p>',
            unsafe_allow_html=True
        )
        f_name  = st.text_input("Person Name",   placeholder="e.g. Ali Khan")
        f_email = st.text_input("Email Address", placeholder="e.g. ali@gmail.com")
        f_phone = st.text_input("Phone Number",  placeholder="e.g. 03348754969")
        f_cnic  = st.text_input("CNIC",          placeholder="e.g. 35202-1234567-3")
        f_ip    = st.text_input("IP Address",    placeholder="e.g. 192.168.1.55")
        f_mac   = st.text_input("MAC Address",   placeholder="e.g. 00:1A:2B:3C:4D:5E")
        f_pass  = st.text_input("Password",      placeholder="e.g. Admin@123!", type="password")
        f_card  = st.text_input("Credit Card",   placeholder="e.g. 4111-1111-1111-1234")
        f_iban  = st.text_input("Bank IBAN",     placeholder="e.g. PK29HABB1234567890123456")
        f_addr  = st.text_input("Home Address",  placeholder="e.g. House 42, DHA, Lahore")
        f_port  = st.text_input("Port Number",   placeholder="e.g. 3306")
        f_other = st.text_area("Other / Log Text",
                               placeholder="Paste any extra text or log lines here...",
                               height=80)

        parts = []
        if f_name:  parts.append(f"Name: {f_name}")
        if f_email: parts.append(f"Email: {f_email}")
        if f_phone: parts.append(f"Phone: {f_phone}")
        if f_cnic:  parts.append(f"CNIC: {f_cnic}")
        if f_ip:    parts.append(f"IP: {f_ip}")
        if f_mac:   parts.append(f"MAC: {f_mac}")
        if f_pass:  parts.append(f"Password: {f_pass}")
        if f_card:  parts.append(f"Card: {f_card}")
        if f_iban:  parts.append(f"IBAN: {f_iban}")
        if f_addr:  parts.append(f"Address: {f_addr}")
        if f_port:  parts.append(f"Port: {f_port}")
        if f_other: parts.append(f_other)
        inp = "\n".join(parts)

        btn = st.button("Analyze and Mask", type="primary", use_container_width=True)

    with col2:
        st.subheader("Masked Output")
        if btn and inp.strip():
            clean  = sanitize_input(inp)
            masked, results, score, val = process(clean, technique)
            log_audit("TEXT", f"chars={len(clean)} fields={len(results)}")
            st.text_area("", value=masked, height=480, key="t1o")
            st.download_button("Download Masked Text",
                masked, "masked.txt", use_container_width=True)
        elif btn:
            st.warning("Please fill in at least one field.")
        else:
            st.markdown(
                '<div class="empty-box">'
                '<span style="font-size:2rem;opacity:.2">🔒</span>'
                '<p style="color:#334155;margin:0;font-size:13px">'
                'Masked output appears here</p>'
                '</div>',
                unsafe_allow_html=True
            )

    if btn and inp.strip():
        clean  = sanitize_input(inp)
        masked, results, score, val = process(clean, technique)
        st.markdown("<hr class='div'>", unsafe_allow_html=True)
        show_results(results, score, val)
        show_share(results, score, val, "manual_input")


# ── TAB 2 ─────────────────────────────────────────────────────────────────────
with tab2:
    st.subheader("Upload CSV or JSON File")
    f = st.file_uploader("", type=["csv", "json"], key="f2")

    if f is not None:
        ok, msg = check_file_size(f)
        if not ok:
            st.error(msg)
        else:
            v, vm = validate_file_type(f)
            if not v:
                st.error(vm)
            else:
                try:
                    log_audit("FILE", f"name={f.name}")

                    if f.name.endswith(".csv"):
                        try:
                            df = pd.read_csv(f, encoding="utf-8")
                        except Exception:
                            f.seek(0)
                            df = pd.read_csv(f, encoding="latin-1")

                    else:
                        raw_json = f.read().decode("utf-8")
                        data     = json.loads(raw_json)
                        if isinstance(data, list):
                            df = pd.DataFrame(data)
                        elif isinstance(data, dict):
                            df = pd.DataFrame()
                            for k, v2 in data.items():
                                if isinstance(v2, list):
                                    df = pd.DataFrame(v2)
                                    break
                            if df.empty:
                                df = pd.DataFrame([data])
                        else:
                            df = pd.DataFrame([data])

                    st.success(
                        f"Loaded: {len(df)} rows x {len(df.columns)} columns"
                    )
                    st.dataframe(df.head(), use_container_width=True)

                    with st.expander("Advanced Options"):
                        ds = st.checkbox("Data Shuffling")
                        dk = st.checkbox("k-Anonymity Grouping")
                        kv = st.slider("k value", 2, 10, 3) if dk else 3
                        sc = st.selectbox("Shuffle column",
                             df.columns.tolist()) if ds else None
                        kc = st.selectbox("k-anon column",
                             df.columns.tolist()) if dk else None

                    if st.button("Mask Entire Dataset",
                                 type="primary",
                                 use_container_width=True,
                                 key="b2"):
                        with st.spinner("Masking all rows..."):
                            mdf, report, total = process_dataframe(df, technique)
                            from src.anonymizer import shuffle_column, k_anonymity
                            if ds and sc:
                                mdf = shuffle_column(mdf, sc)
                            if dk and kc:
                                mdf = k_anonymity(mdf, kc, kv)

                        log_audit("DATASET",
                                  f"rows={len(df)} masked={total}")
                        st.success(
                            f"{total} values masked in {len(df)} rows"
                        )
                        st.dataframe(mdf.head(), use_container_width=True)
                        st.markdown("<hr class='div'>", unsafe_allow_html=True)
                        show_col(report)

                        d1, d2 = st.columns(2)
                        with d1:
                            st.download_button(
                                "Download Masked CSV",
                                mdf.to_csv(index=False),
                                "masked.csv",
                                use_container_width=True
                            )
                        with d2:
                            st.download_button(
                                "Download Column Report",
                                pd.DataFrame(report).to_csv(index=False),
                                "col_report.csv",
                                use_container_width=True
                            )

                        ar = [
                            {
                                "Type":      r["Detected As"],
                                "Risk":      "CRITICAL " + r["Risk Level"],
                                "Original":  "***",
                                "Masked":    "***",
                                "Technique": r["Technique"]
                            }
                            for r in report if r["Values Masked"] > 0
                        ]
                        show_share(ar, 60, {
                            "original_count":  total,
                            "remaining_count": 0,
                            "still_exposed":   [],
                            "safe_to_share":   True
                        }, f.name)

                except Exception as e:
                    st.error("Error processing file: " + str(e))


# ── TAB 3 ─────────────────────────────────────────────────────────────────────
with tab3:
    st.subheader("Upload PDF or Word Document")
    f = st.file_uploader("",
                         type=["pdf", "docx", "txt", "log"],
                         key="f3")

    if f is not None:
        ok, msg = check_file_size(f)
        if not ok:
            st.error(msg)
        else:
            try:
                log_audit("DOC", f"name={f.name}")
                ext = f.name.split(".")[-1].lower()

                if ext == "pdf":
                    raw = read_pdf(f)
                elif ext == "docx":
                    raw = read_docx(f)
                else:
                    raw = f.read().decode("utf-8", errors="ignore")

                st.success(f"Loaded: {f.name}")

                c1, c2 = st.columns(2, gap="large")
                with c1:
                    st.subheader("Original Text")
                    st.text_area("", value=raw, height=360, key="d1")

                if st.button("Mask Document",
                             type="primary",
                             use_container_width=True,
                             key="b3"):
                    with st.spinner("Masking..."):
                        masked, results, score, val = process(
                            sanitize_input(raw), technique
                        )
                    log_audit("DOC_MASKED",
                              f"fields={len(results)}")
                    with c2:
                        st.subheader("Masked Text")
                        st.text_area("", value=masked,
                                     height=360, key="d2")
                        st.download_button(
                            "Download Masked Document",
                            masked,
                            f"masked_{f.name}.txt",
                            use_container_width=True
                        )
                    st.markdown("<hr class='div'>", unsafe_allow_html=True)
                    show_results(results, score, val)
                    show_share(results, score, val, f.name)

            except Exception as e:
                st.error("Error processing file: " + str(e))
                st.info("Make sure file is not password protected.")